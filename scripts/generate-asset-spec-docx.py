# -*- coding: utf-8 -*-
"""Generate Rhythm Guard image asset specification Word document."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "RhythmGuard-Image-Asset-Spec.docx"
OUT_CN = ROOT / "docs" / "RhythmGuard-\u56fe\u7247\u7d20\u6750\u9700\u6c42\u6c47\u603b.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        from docx.oxml import OxmlElement

        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "D9E2F3") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    title = doc.add_heading("Rhythm Guard 图片素材需求汇总", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("项目：Rhythm Guard（节奏守卫）\n版本：v2.1.0 | 棋盘 10×10 | 主场景 960×640 | 节奏场景 1080×1920")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph(
        "本文档汇总当前游戏所需全部图片资源，包含推荐像素尺寸、游戏内显示尺寸及用途说明。"
        "标注「程序绘制」的项目目前由 Phaser 几何图形实现，可暂不制作；若需美术替换则按表中尺寸提供。"
    )

    doc.add_heading("一、全局规范", level=1)
    add_table(
        doc,
        ["项目", "数值"],
        [
            ["主场景分辨率", "960 × 640（菜单 / 部署 / 塔防 / 结算 / Meta）"],
            ["节奏场景分辨率", "1080 × 1920（竖屏）"],
            ["棋盘规格", "10×10，单格 48×48 px（内容区 44×44）"],
            ["城堡核心区域", "占 2×2 格 ≈ 96×96 px，碰撞圆半径 35 px"],
            ["推荐导出策略", "游戏内使用 @1x；原画建议按 @2x 绘制后缩放"],
            ["文件格式", "PNG（透明底）为主；大背景可用 JPG/WebP"],
            ["目录建议", "public/assets/sprites/{类别}/{名称}@2x.png"],
        ],
    )

    doc.add_heading("二、品牌与通用 UI（960×640）", level=1)
    add_table(
        doc,
        ["资源 ID", "内容", "画布尺寸 (@1x)", "游戏内显示", "用途"],
        [
            ["ui/bg_main", "主场景深色背景", "960×640", "全屏", "菜单/部署/塔防/结算/Meta 底图"],
            ["ui/bg_rhythm", "节奏场景暖色背景", "1080×1920", "全屏", "节奏阶段底图"],
            ["ui/btn_normal", "按钮九宫格", "200×88（@2x: 400×176）", "高 44，宽随文字", "所有文字按钮"],
            ["ui/btn_hover", "按钮悬停态", "同上", "同上", "按钮交互反馈"],
            ["ui/panel_meta", "Meta 升级行背景", "800×52", "约 800×52", "Meta 每行装饰（可选）"],
            ["ui/icon_star", "星星图标", "32×32", "16～36 px", "关卡星星、总星数"],
            ["ui/icon_gold", "金币图标", "24×24", "HUD 行内", "塔防金币（现用 emoji）"],
            ["brand/logo_title", "游戏 Logo", "480×80", "居中置顶", "主菜单标题"],
            ["brand/favicon", "站点图标", "已有 favicon.svg", "32×32", "浏览器标签 / PWA"],
        ],
    )
    doc.add_paragraph("按钮参考色：紫 #7b4fff、绿 #27ae60、红 #e74c3c、灰 #444444")

    doc.add_heading("三、节奏阶段（1080×1920）", level=1)
    doc.add_heading("3.1 核心判定区", level=2)
    add_table(
        doc,
        ["资源 ID", "内容", "画布 (@1x)", "游戏内显示", "用途"],
        [
            ["rhythm/face_body", "中心节奏圆盘", "540×540", "半径 270 px", "TargetFace 主体，随音符变色"],
            ["rhythm/face_eye", "眼睛（单眼）", "104×146", "椭圆约 52×73", "左右眼各一"],
            ["rhythm/guide_ring", "浅黄引导外圈", "580×580（环宽 ~22）", "外半径 270 + 描边", "对齐参考环"],
            ["rhythm/ring_yellow", "缩圈描边（黄）", "环宽 22 px", "动态缩放", "ShrinkRing 黄圈"],
            ["rhythm/ring_blue", "缩圈描边（蓝）", "同上", "同上", "蓝圈长按"],
            ["rhythm/ring_red", "缩圈描边（红）", "同上", "同上", "红圈连击"],
            ["rhythm/hold_progress", "长按进度弧", "580×580 环形", "半径 ~276", "HoldProgressRing"],
            ["rhythm/hit_fx_perfect", "PERFECT 光晕", "128×128", "动态缩放", "判定特效"],
            ["rhythm/hit_fx_great", "GREAT 光晕", "128×128", "同上", "判定特效"],
            ["rhythm/hit_fx_good", "GOOD 光晕", "128×128", "同上", "判定特效"],
            ["rhythm/hit_fx_miss", "MISS 光晕", "128×128", "同上", "判定特效"],
        ],
    )
    doc.add_paragraph(
        "关键尺寸：中心圆外径 540 px；缩圈起始外径约 1690 px（建议程序绘制）；"
        "点击热区半径 350 px；顶部进度条 1000×12 px，紫色 #7b4fff。"
    )

    doc.add_heading("3.2 节奏装饰（可选）", level=2)
    add_table(
        doc,
        ["资源 ID", "内容", "画布尺寸", "用途"],
        [
            ["rhythm/deco_toy_01～07", "背景装饰物", "各 84×84", "替换 7 个浮动 emoji 装饰"],
            ["rhythm/progress_bar_bg", "进度条底", "1000×12", "节奏倒计时条背景"],
            ["rhythm/progress_bar_fill", "进度条填充", "1000×12", "可九宫格横向拉伸"],
            ["rhythm/overlay_end", "节奏结束遮罩", "1080×1920", "半透明结算层"],
        ],
    )

    doc.add_heading("3.3 节奏文字（可用位图字或系统字体）", level=2)
    add_table(
        doc,
        ["内容", "字号", "用途"],
        [
            ["PERFECT / GREAT / GOOD / MISS", "40 px", "判定飘字"],
            ["Score / Combo", "26 px", "左上角 HUD"],
            ["关卡名", "32 px", "顶部标题"],
        ],
    )

    doc.add_heading("四、塔 / 道具（蘑菇守卫）", level=1)
    doc.add_paragraph("4 族 × 4 Tier = 16 张独立图标（或 4 张底座 + 4 张阶数角标）。")
    add_table(
        doc,
        ["资源 ID", "名称", "画布 (@1x)", "游戏内显示", "用途"],
        [
            ["tower/kick_t1", "幼菇拳手", "96×96", "部署 52×52 / 防守 r=22", "Kick T1 近战"],
            ["tower/kick_t2", "壮菇卫士", "96×96", "同上", "Kick T2"],
            ["tower/kick_t3", "铁菇冲锋", "96×96", "同上", "Kick T3"],
            ["tower/kick_t4", "王菇禁卫", "96×96", "同上", "Kick T4"],
            ["tower/snare_t1", "毒芽孢", "96×96", "同上", "Snare T1 溅射（蓝）"],
            ["tower/snare_t2", "爆裂孢", "96×96", "同上", "Snare T2"],
            ["tower/snare_t3", "腐毒菇", "96×96", "同上", "Snare T3"],
            ["tower/snare_t4", "瘟疫领主", "96×96", "同上", "Snare T4"],
            ["tower/hihat_t1", "静电菇", "96×96", "同上", "Hihat T1 连锁（黄）"],
            ["tower/hihat_t2", "弧光菇", "96×96", "同上", "Hihat T2"],
            ["tower/hihat_t3", "雷暴菇", "96×96", "同上", "Hihat T3"],
            ["tower/hihat_t4", "天罚菇", "96×96", "同上", "Hihat T4"],
            ["tower/crash_t1", "碎石护符", "96×96", "同上", "Crash T1 护盾（紫）"],
            ["tower/crash_t2", "岩盾结界", "96×96", "同上", "Crash T2"],
            ["tower/crash_t3", "堡垒蘑菇", "96×96", "同上", "Crash T3"],
            ["tower/crash_t4", "不朽菌墙", "96×96", "同上", "Crash T4"],
        ],
    )
    add_table(
        doc,
        ["资源 ID", "尺寸", "用途"],
        [
            ["tower/badge_tier_1～4", "24×24", "阶数角标，替代 T1～T4 文字"],
            ["tower/badge_melee/aoe/chain/shield", "20×20", "攻击类型小标（近/爆/链/盾）"],
            ["tower/range_ring", "264×264", "最大射程 132 px 半透明圈（可选）"],
        ],
        header_fill="E2EFDA",
    )

    doc.add_heading("五、敌人（入侵英雄）10 种", level=1)
    doc.add_paragraph("显示尺寸 = JSON 中 size（半径）；精灵图建议直径 + 上下留白。")
    add_table(
        doc,
        ["资源 ID", "名称", "游戏内半径", "推荐画布 (@1x)", "备注"],
        [
            ["enemy/goblin", "见习勇者", "11 px", "44×56", "人海，轻甲"],
            ["enemy/grunt", "王国步兵", "16 px", "56×68", "标准步兵"],
            ["enemy/fast", "疾风游侠", "12 px", "48×60", "高速"],
            ["enemy/assassin", "暗影剑客", "13 px", "50×62", "精英快攻"],
            ["enemy/shielder", "圣盾骑士", "18 px", "60×74", "护甲，银边"],
            ["enemy/tank", "重装圣骑士", "22 px", "68×82", "重甲"],
            ["enemy/healer", "祝福牧师", "15 px", "54×70", "治疗光环"],
            ["enemy/flyer", "狮鹫骑士", "14 px", "52×52", "飞行，三角/菱形"],
            ["enemy/brute", "攻城巨像", "26 px", "76×90", "精英"],
            ["enemy/warlord", "传奇英雄 Boss", "32 px", "96×110", "Boss 金边"],
        ],
    )
    add_table(
        doc,
        ["资源 ID", "尺寸", "用途"],
        [
            ["enemy/badge_flying/armor/heal/boss", "32×16 或 40×16", "单位标签"],
            ["enemy/hp_bar_bg / hp_bar_fill", "34×4", "血条底/填充"],
            ["enemy/aura_heal", "170×170", "牧师治疗圈（半径 85）"],
        ],
        header_fill="FCE4D6",
    )

    doc.add_heading("六、棋盘与城堡（部署 + 塔防）", level=1)
    add_table(
        doc,
        ["资源 ID", "内容", "画布 (@1x)", "游戏内显示", "用途"],
        [
            ["grid/cell_normal", "普通格", "48×48", "单格", "可放置格"],
            ["grid/cell_core", "城堡核心格", "48×48", "2×2 四格", "紫底金边核心"],
            ["grid/cell_hover", "放置高亮", "48×48", "单格", "拖拽悬停（可选）"],
            ["grid/cell_invalid", "不可放置", "48×48", "单格", "核心区提示（可选）"],
            ["grid/board_frame", "棋盘外框", "480×480", "10×10 整体", "装饰边框（可选）"],
            ["castle/core", "蘑菇城堡主体", "96×96 或 128×128", "圆 r=35", "城堡视觉"],
            ["castle/core_shield", "护盾脉冲环", "128×128", "动态缩放", "岩盾菇/受击反馈"],
            ["castle/hp_bar_bg", "城堡血条底", "300×16", "底部 HUD", "塔防 HP 条"],
            ["castle/hp_bar_fill", "城堡血条填充", "300×16", "动态宽度", "HP 填充"],
        ],
    )

    doc.add_heading("七、战斗特效（可选，现程序绘制）", level=1)
    add_table(
        doc,
        ["资源 ID", "内容", "推荐尺寸", "用途"],
        [
            ["fx/melee_slash", "近战斩击", "128×8 或 4 帧序列", "Kick 攻击"],
            ["fx/aoe_burst", "爆炸圈", "64～154", "Snare 溅射（最大 AOE 半径 77）"],
            ["fx/chain_bolt", "闪电链接", "128×128 序列 3～6 帧", "Hihat 连锁"],
            ["fx/shield_absorb", "护盾吸收", "96×96", "护盾抵挡"],
            ["fx/core_hit", "城堡受击", "96×96", "核心闪红"],
            ["fx/spawn_flash", "刷怪闪光", "48×48", "敌人出生"],
            ["fx/damage_number_bg", "伤害飘字底", "64×24", "可选"],
            ["fx/enemy_death", "敌人死亡", "48×48 序列 4 帧", "替代缩小消失"],
        ],
    )

    doc.add_heading("八、关卡与章节", level=1)
    add_table(
        doc,
        ["资源 ID", "内容", "尺寸", "用途"],
        [
            ["chapter/ch1_icon", "蘑菇王国", "64×64", "第 1 章（绿 #27ae60）"],
            ["chapter/ch2_icon", "暗影森林", "64×64", "第 2 章（#2c3e50）"],
            ["chapter/ch3_icon", "雷霆高地", "64×64", "第 3 章（#8e44ad）"],
            ["level/icon_locked", "关卡锁", "32×32", "未解锁"],
            ["level/icon_cleared", "关卡通", "32×32", "已通关（可选）"],
        ],
    )
    doc.add_paragraph("共 8 关，可用统一按钮底图 + 文字，不强制 8 张独立图。")

    doc.add_heading("九、Meta 升级图标", level=1)
    add_table(
        doc,
        ["资源 ID", "内容", "尺寸", "用途"],
        [
            ["meta/icon_timing", "节奏感知", "48×48", "判定窗口加成"],
            ["meta/icon_merge", "合成大师", "48×48", "合成次数加成"],
            ["meta/icon_fortress", "堡垒强化", "48×48", "城堡 HP 加成"],
            ["meta/icon_fortune", "幸运律动", "48×48", "掉落加成"],
        ],
    )

    doc.add_heading("十、资源数量汇总", level=1)
    add_table(
        doc,
        ["类别", "必做（换美术最低）", "完整（含可选）"],
        [
            ["通用 UI / 背景", "4", "10"],
            ["节奏核心", "8", "18"],
            ["塔 / 道具", "16", "24"],
            ["敌人", "10", "18"],
            ["棋盘 / 城堡", "4", "10"],
            ["战斗特效", "0（可程序）", "8"],
            ["章节 / Meta", "0（可文字）", "11"],
            ["合计", "约 42 张", "约 99 张"],
        ],
        header_fill="FFF2CC",
    )

    doc.add_heading("十一、目录结构建议", level=1)
    code = doc.add_paragraph()
    code.add_run(
        "public/assets/\n"
        "├── sprites/\n"
        "│   ├── ui/\n"
        "│   ├── rhythm/\n"
        "│   ├── tower/\n"
        "│   ├── enemy/\n"
        "│   ├── grid/\n"
        "│   ├── castle/\n"
        "│   ├── fx/\n"
        "│   ├── chapter/\n"
        "│   └── meta/\n"
        "├── audio/          （BGM / 音效）\n"
        "└── charts/         （曲谱 JSON）"
    ).font.name = "Consolas"

    doc.add_heading("十二、制作优先级建议", level=1)
    add_table(
        doc,
        ["优先级", "内容", "张数约"],
        [
            ["P0", "tower/* 16 张 + enemy/* 10 张 + castle/core + grid/cell_*", "约 27～31"],
            ["P1", "rhythm/face_body + 三环描边 + 节奏背景", "约 8～12"],
            ["P2", "UI 按钮、特效序列帧、章节/Meta 图标", "其余"],
        ],
        header_fill="DDEBF7",
    )

    doc.add_heading("十三、AI 批量生成额度参考（Cursor）", level=1)
    add_table(
        doc,
        ["批次", "张数", "图片 API 费用（估）", "Pro $20 池占比（估）"],
        [
            ["试水批", "5～8", "$0.7～2", "3%～10%"],
            ["P0 核心", "约 27", "$3.5～7", "18%～35%"],
            ["最低可换美术", "约 42", "$5～10", "25%～50%"],
            ["较完整", "约 99", "$13～24", "可能超出月池"],
        ],
    )
    doc.add_paragraph(
        "说明：生图按 API 计费（约 $0.13/张 1K～2K）；含重绘时按 1.2～1.5 倍估算；"
        "Agent 对话另计 token 费用。实际消耗以 Cursor Settings → Account → Usage 为准。"
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("生成日期：2026-06-05 | Rhythm Guard v2.1.0").font.size = Pt(9)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    doc.save(OUT)
    doc.save(OUT_CN)
    print(f"Saved: {OUT}")
    print(f"Saved: {OUT_CN}")


if __name__ == "__main__":
    main()
